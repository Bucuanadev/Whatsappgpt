import os
import uuid
import hashlib
from typing import Dict, List, Optional
from dataclasses import dataclass
from datetime import datetime
import json

# Import for embeddings and vector operations
try:
    from langchain_openai import OpenAIEmbeddings
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.docstore.document import Document
    from pinecone import Pinecone
except ImportError as e:
    print(f"Warning: Some dependencies not available: {e}")

@dataclass
class DocumentChunk:
    """Data class for document chunks"""
    chunk_id: str
    client_id: str
    content: str
    metadata: Dict
    embedding: Optional[List[float]] = None
    created_at: str = None
    
    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now().isoformat()

class VectorStoreManager:
    """Manager for vector store operations using Pinecone"""
    
    def __init__(self, api_key: str = None, environment: str = None, index_name: str = None):
        self.api_key = api_key or os.getenv('PINECONE_API_KEY')
        self.environment = environment or os.getenv('PINECONE_ENVIRONMENT', 'us-east-1-aws')
        self.index_name = index_name or os.getenv('PINECONE_INDEX_NAME', 'whatsapp-gpt')
        
        # Initialize OpenAI embeddings
        self.embeddings = OpenAIEmbeddings(
            openai_api_key=os.getenv('OPENAI_API_KEY')
        )
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Initialize Pinecone
        self.index = None
        self._initialize_pinecone()
    
    def _initialize_pinecone(self):
        """Initialize Pinecone connection"""
        try:
            if self.api_key:
                # Initialize Pinecone client
                pc = Pinecone(api_key=self.api_key)
                
                # Check if index exists, create if not
                existing_indexes = pc.list_indexes()
                index_names = [idx['name'] for idx in existing_indexes]
                
                if self.index_name not in index_names:
                    pc.create_index(
                        name=self.index_name,
                        dimension=1536,  # OpenAI embedding dimension
                        metric='cosine'
                    )
                
                self.index = pc.Index(self.index_name)
            else:
                print("Warning: Pinecone API key not provided")
        except Exception as e:
            print(f"Error initializing Pinecone: {str(e)}")
    
    def create_namespace(self, client_id: str) -> str:
        """
        Create isolated namespace for client's knowledge base
        
        Args:
            client_id: Unique identifier for the client
            
        Returns:
            Namespace identifier
        """
        # Use client_id as namespace
        namespace = f"client_{client_id}"
        
        return namespace
    
    def process_text_content(self, client_id: str, content: str, metadata: Dict = None) -> Dict:
        """
        Process text content and add to vector store
        
        Args:
            client_id: Unique identifier for the client
            content: Text content to process
            metadata: Additional metadata for the content
            
        Returns:
            Dictionary with processing result
        """
        try:
            if metadata is None:
                metadata = {}
            
            # Split text into chunks
            documents = [Document(page_content=content, metadata=metadata)]
            text_chunks = self.text_splitter.split_documents(documents)
            
            # Create namespace
            namespace = self.create_namespace(client_id)
            
            # Process each chunk
            processed_chunks = []
            vectors_to_upsert = []
            
            for i, chunk in enumerate(text_chunks):
                # Generate unique ID for chunk
                chunk_id = f"{client_id}_{hashlib.md5(chunk.page_content.encode()).hexdigest()}_{i}"
                
                # Generate embedding
                embedding = self.embeddings.embed_query(chunk.page_content)
                
                # Prepare metadata
                chunk_metadata = {
                    **chunk.metadata,
                    'client_id': client_id,
                    'chunk_index': i,
                    'content_type': 'text',
                    'created_at': datetime.now().isoformat()
                }
                
                # Create document chunk
                doc_chunk = DocumentChunk(
                    chunk_id=chunk_id,
                    client_id=client_id,
                    content=chunk.page_content,
                    metadata=chunk_metadata,
                    embedding=embedding
                )
                
                processed_chunks.append(doc_chunk)
                
                # Prepare vector for upsert
                vectors_to_upsert.append({
                    'id': chunk_id,
                    'values': embedding,
                    'metadata': {
                        **chunk_metadata,
                        'content': chunk.page_content[:1000]  # Limit content size in metadata
                    }
                })
            
            # Upsert vectors to Pinecone
            if self.index and vectors_to_upsert:
                self.index.upsert(
                    vectors=vectors_to_upsert,
                    namespace=namespace
                )
            
            return {
                "success": True,
                "chunks_processed": len(processed_chunks),
                "namespace": namespace,
                "message": "Text content processed successfully"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Exception processing text content: {str(e)}"
            }
    
    def process_document_file(self, client_id: str, file_path: str, file_type: str = None) -> Dict:
        """
        Process document file and add to vector store
        
        Args:
            client_id: Unique identifier for the client
            file_path: Path to the document file
            file_type: Type of document (pdf, txt, docx, etc.)
            
        Returns:
            Dictionary with processing result
        """
        try:
            # Extract text from file based on type
            content = self._extract_text_from_file(file_path, file_type)
            
            if not content:
                return {
                    "success": False,
                    "error": "Could not extract text from file"
                }
            
            # Process extracted content
            metadata = {
                'source_file': os.path.basename(file_path),
                'file_type': file_type or 'unknown',
                'source_type': 'document'
            }
            
            return self.process_text_content(client_id, content, metadata)
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Exception processing document: {str(e)}"
            }
    
    def _extract_text_from_file(self, file_path: str, file_type: str = None) -> str:
        """
        Extract text content from various file types
        
        Args:
            file_path: Path to the file
            file_type: Type of file
            
        Returns:
            Extracted text content
        """
        try:
            if not os.path.exists(file_path):
                return ""
            
            # Determine file type if not provided
            if not file_type:
                _, ext = os.path.splitext(file_path)
                file_type = ext.lower().lstrip('.')
            
            # Extract text based on file type
            if file_type in ['txt', 'md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif file_type == 'pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except ImportError:
                    print("PyPDF2 not available for PDF processing")
                    return ""
            
            elif file_type in ['doc', 'docx']:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    return text
                except ImportError:
                    print("python-docx not available for DOCX processing")
                    return ""
            
            else:
                # Try to read as text file
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        return f.read()
                except:
                    return ""
                    
        except Exception as e:
            print(f"Error extracting text from file: {str(e)}")
            return ""
    
    def query_knowledge_base(self, client_id: str, query: str, top_k: int = 5) -> Dict:
        """
        Query client's knowledge base for relevant information
        
        Args:
            client_id: Unique identifier for the client
            query: Search query
            top_k: Number of top results to return
            
        Returns:
            Dictionary with query results
        """
        try:
            if not self.index:
                return {
                    "success": False,
                    "error": "Vector store not initialized"
                }
            
            # Generate query embedding
            query_embedding = self.embeddings.embed_query(query)
            
            # Create namespace
            namespace = self.create_namespace(client_id)
            
            # Query Pinecone
            results = self.index.query(
                vector=query_embedding,
                top_k=top_k,
                namespace=namespace,
                include_metadata=True
            )
            
            # Process results
            relevant_chunks = []
            for match in results.matches:
                relevant_chunks.append({
                    'content': match.metadata.get('content', ''),
                    'score': match.score,
                    'metadata': match.metadata
                })
            
            # Combine relevant content
            context = "\n\n".join([chunk['content'] for chunk in relevant_chunks])
            
            return {
                "success": True,
                "context": context,
                "chunks": relevant_chunks,
                "total_results": len(relevant_chunks)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Exception querying knowledge base: {str(e)}"
            }
    
    def delete_client_data(self, client_id: str) -> Dict:
        """
        Delete all data for a client
        
        Args:
            client_id: Unique identifier for the client
            
        Returns:
            Dictionary with deletion result
        """
        try:
            if not self.index:
                return {
                    "success": False,
                    "error": "Vector store not initialized"
                }
            
            namespace = self.create_namespace(client_id)
            
            # Delete all vectors in the namespace
            self.index.delete(delete_all=True, namespace=namespace)
            
            return {
                "success": True,
                "message": "Client data deleted successfully"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Exception deleting client data: {str(e)}"
            }
    
    def get_client_stats(self, client_id: str) -> Dict:
        """
        Get statistics for client's knowledge base
        
        Args:
            client_id: Unique identifier for the client
            
        Returns:
            Dictionary with statistics
        """
        try:
            if not self.index:
                return {
                    "success": False,
                    "error": "Vector store not initialized"
                }
            
            namespace = self.create_namespace(client_id)
            
            # Get index stats
            stats = self.index.describe_index_stats()
            namespace_stats = stats.namespaces.get(namespace, {})
            
            return {
                "success": True,
                "total_vectors": namespace_stats.get('vector_count', 0),
                "namespace": namespace,
                "dimension": stats.dimension
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Exception getting client stats: {str(e)}"
            }

